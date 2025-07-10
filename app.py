# app.py
"""
Gradio 主入口，整合 rag_core 的各个模块，提供上传文档、提问问答、
向量数据库查看与管理等功能的可视化界面。
"""

import os
import fitz
from pathlib import Path
import shutil
import docx  # 读取 Word 文件

import gradio as gr

from rag_core.doc_processor import DocumentProcessor
from rag_core.embedder import EmbeddingModel
from rag_core.vector_store import VectorStore
from rag_core.generator import AnswerGenerator

import os
import fitz
from pathlib import Path
# ...

import gradio as gr

# 🔽 读取外部 CSS 文件
with open("style.css", "r", encoding="utf-8") as f:
    css = f.read()


# 初始化 RAG 组件
API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-de353379ea494b98b197dce3b8ce5391")  # 你的API KEY
processor = DocumentProcessor()
embedder = EmbeddingModel()
vector_store = VectorStore(embedder.embedding_dim)
generator = AnswerGenerator(API_KEY)

def upload_document(file):
    if not file:
        return "❌ 请上传有效的文件。"

    ext = Path(file.name).suffix.lower()
    if ext == ".pdf":
        sub_dir = "pdf"
    elif ext == ".docx":
        sub_dir = "docx"
    elif ext == ".txt":
        sub_dir = "txt"
    else:
        return "❌ 不支持的文件类型（仅支持 PDF / DOCX / TXT）"

    original_filename = Path(file.name).name

    # 判断是否已上传过相同文档
    if any(doc.source == original_filename for doc in vector_store.documents):
        return f"⚠️ 检测到已存在相同名称的文档：{original_filename}，已跳过上传。"

    save_dir = os.path.join("raw_data", sub_dir)
    os.makedirs(save_dir, exist_ok=True)
    save_path = os.path.join(save_dir, original_filename)
    shutil.copyfile(file.name, save_path)

    # 读取文件内容
    if ext == ".pdf":
        doc = fitz.open(file.name)
        text = "\n".join([page.get_text() for page in doc])
    elif ext == ".docx":
        text = "\n".join([para.text for para in docx.Document(file.name).paragraphs])
    elif ext == ".txt":
        with open(file.name, "r", encoding="utf-8") as f:
            text = f.read()

    docs = processor.process_document(text, title=original_filename, source=original_filename)
    embeddings = embedder.encode_texts([doc.content for doc in docs])
    vector_store.add_documents(docs, embeddings)

    return f"✅ 成功上传并解析 {ext[1:].upper()} 文件，共 {len(docs)} 个片段"


# 在 app.py 中添加（可放在 upload_document 函数下方）
from pathlib import Path
import os
import fitz
import docx

def sync_raw_data_to_db(vector_store):
    """
    检查 raw_data 目录中所有文件，将未编码到向量数据库的文件编码并加入数据库
    返回处理结果日志
    """
    # 1. 扫描 raw_data 下所有支持的文件（PDF、DOCX、TXT）
    supported_ext = {".pdf", ".docx", ".txt"}
    raw_files = []
    for root, _, files in os.walk("raw_data"):
        for file in files:
            if Path(file).suffix.lower() in supported_ext:
                raw_files.append(os.path.join(root, file))

    # 2. 获取已处理文件名（直接从 vector_store.documents 读取）
    processed_sources = set()
    for doc in vector_store.documents:
        if hasattr(doc, "source"):
            processed_sources.add(doc.source)

    # 3. 筛选未处理文件
    unprocessed_files = []
    for file_path in raw_files:
        file_name = Path(file_path).name
        if file_name not in processed_sources:
            unprocessed_files.append(file_path)

    if not unprocessed_files:
        return "✅ 所有 raw_data 中的文件已编码到向量数据库，无需处理"

    # 4. 处理未编码文件
    log = [f"发现 {len(unprocessed_files)} 个未处理文件，开始编码..."]
    for file_path in unprocessed_files:
        try:
            ext = Path(file_path).suffix.lower()
            if ext == ".pdf":
                doc = fitz.open(file_path)
                text = "\n".join([page.get_text() for page in doc])
            elif ext == ".docx":
                text = "\n".join([para.text for para in docx.Document(file_path).paragraphs])
            elif ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                log.append(f"❌ 跳过不支持的文件：{file_path}")
                continue

            file_name = Path(file_path).name
            docs = processor.process_document(text, title=file_name, source=file_name)
            embeddings = embedder.encode_texts([doc.content for doc in docs])
            vector_store.add_documents(docs, embeddings)
            log.append(f"✅ 成功编码：{file_name}（{len(docs)} 个片段）")
        except Exception as e:
            log.append(f"❌ 处理 {file_path} 失败：{str(e)}")

    return "\n".join(log)


import re
import markdown

import markdown
import html

def qa_interface(question):
    # 获取向量 + 检索结果 + 基础与增强回答
    q_embed = embedder.encode_query(question)
    top_docs = vector_store.search(q_embed)
    base_answer = generator.generate_base_answer(question)
    enhanced = generator.generate_enhanced_answer(question, [doc for doc, _ in top_docs], base_answer)

    # 置信度
    confidence = sum(score for _, score in top_docs) / (len(top_docs) or 1)

    # 文档片段：做安全转义和换行美化
    result_docs = "\n".join([f"[{i+1}] {html.escape(doc.content[:100])}..." for i, (doc, _) in enumerate(top_docs)])
    docs_html = result_docs.replace('\n', '<br>')

    # Markdown -> HTML（支持表格、代码块）
    enhanced_html = markdown.markdown(
        enhanced,
        extensions=["tables", "fenced_code"]
    )

    # 构造完整 HTML 页面片段
    html_response = f"""
    <div id="custom-html-answer" style="font-family: 'Segoe UI', sans-serif; line-height: 1.6;">
        <div>
            <h3>【回复】</h3>
            <div>{enhanced_html}</div>
        </div>

        <div style="margin-top: 1.5em;">
            <h3>【来源文档片段】</h3>
            <div style="font-size: 0.95em; color: #555;">{docs_html}</div>
        </div>

        <div style="margin-top: 1em;">
            <strong>置信度:</strong> {confidence:.2f}
        </div>
    </div>
    """

    return html_response



def list_docs():
    # 先同步 raw_data 到数据库，再返回最新的文档列表
    sync_log = sync_raw_data_to_db(vector_store)
    latest_docs = "\n".join(vector_store.list_documents())
    # 返回同步日志和最新文档列表（用元组包装多个返回值）
    return sync_log, latest_docs

def search_by_keyword(keyword):
    return "\n".join(vector_store.search_by_keyword(keyword))

def delete_by_index(index):
    try:
        return vector_store.delete_document_by_index(int(index))
    except:
        return "❌ 输入的编号无效，请输入整数。"

def batch_delete_by_indices(indices_str: str):
    import re

    indices = set()
    try:
        parts = [p.strip() for p in indices_str.split(",")]
        for part in parts:
            if re.match(r"^\d+-\d+$", part):
                start, end = map(int, part.split("-"))
                if start > end:
                    return "❌ 区间格式错误，起始索引应小于等于结束索引"
                indices.update(range(start, end + 1))
            elif part.isdigit():
                indices.add(int(part))
            else:
                return f"❌ 输入格式错误，不能解析部分：'{part}'"
    except Exception as e:
        return f"❌ 解析输入时发生错误: {e}"

    results = []
    for index in sorted(indices, reverse=True):
        results.append(vector_store.delete_document_by_index(index))

    return "\n".join(results)

def show_thinking():
    return "<p style='font-size:28px; font-weight:bold; color:#666;'>🤔 结合数据库搜索，深度思考中......</p>"




with gr.Blocks(css=css) as app:
    with gr.Tabs():
        # ✅ 问答页面
        with gr.TabItem("💬 问答页面"):
            with gr.Column(scale=3):
                # 🔹 回复展示区域（滚动框）
                with gr.Group(elem_classes="fixed-box"):
                    answer_output = gr.HTML(elem_id="custom-html-answer")

                # 🔹 输入框 + 提问按钮 横向排列
                with gr.Row():
                    with gr.Column(scale=9):
                        question_input = gr.Textbox(lines=2,max_lines=2,label="❓ 请输入问题", elem_id="question-box")
                    with gr.Column(scale=1):
                        ask_btn = gr.Button("🚀 发送", elem_id="ask-button", elem_classes="qa-button")


            # 🔹 问答逻辑绑定
            ask_btn.click(show_thinking, inputs=None, outputs=answer_output)
            question_input.submit(show_thinking, inputs=None, outputs=answer_output)

            ask_btn.click(qa_interface, inputs=question_input, outputs=answer_output, queue=True)
            question_input.submit(qa_interface, inputs=question_input, outputs=answer_output, queue=True)

        # ✅ 数据库管理页面
        with gr.TabItem("📁 数据库检索"):
            with gr.Row():
                with gr.Column(scale=1):
                    gr.Markdown("### 📋 文档上传")
                    pdf_input = gr.File(label="📎 上传文档", file_types=[".pdf", ".docx", ".txt"])
                    upload_btn = gr.Button("📤 确认上传", elem_classes="db-button")
                    upload_feedback = gr.Textbox(lines=7,max_lines=7,label="📩 上传反馈", interactive=False)

                with gr.Column(scale=3):
                    gr.Markdown("### 📋 文档列表")
                    sync_log_output = gr.Textbox(lines=3,label="同步日志", max_lines=3, interactive=False)
                    list_btn = gr.Button("🔄 刷新列表", elem_classes="db-button")
                    list_output = gr.Textbox(lines=13,label="📚 当前文档片段", max_lines=13, interactive=False)  # 建议增加行数

                with gr.Column(scale=3):
                    gr.Markdown("### 🔍 数据库关键词搜索")
                    search_input = gr.Textbox(lines=3,max_lines=3,label="📝 请输入要检索的文档关键词：")
                    search_btn = gr.Button("🔍 搜索", elem_classes="db-button")
                    search_output = gr.Textbox(lines=13,max_lines=13,label="📄 检索结果", interactive=False)
        
        with gr.TabItem("📁 数据库删除"):
            with gr.Column(scale=1):
                gr.Markdown("### 🗑️ 删除文档片段")
                delete_input = gr.Textbox(lines=3,max_lines=3,label="✂️ 删除编号（支持单个、逗号分隔、区间，如：3、1,4、2-5）")
                delete_btn = gr.Button("🗑️ 删除", elem_classes="db-button")
                delete_output = gr.Textbox(lines=13,max_lines=13,label="🧾 删除结果", interactive=False)

    # ✅ 功能绑定
    upload_btn.click(upload_document, inputs=pdf_input, outputs=upload_feedback)
    # 功能绑定部分修改为：
    list_btn.click(
        fn=list_docs,
        inputs=None,
        outputs=[sync_log_output, list_output]  # 两个输出分别对应 sync_log 和 latest_docs
    )
    search_btn.click(search_by_keyword, inputs=search_input, outputs=search_output)
    delete_btn.click(batch_delete_by_indices, inputs=delete_input, outputs=delete_output)


if __name__ == "__main__":
    app.launch()
